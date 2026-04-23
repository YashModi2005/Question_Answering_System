from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def set_page_border(doc):
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)
        sect_pr.append(pg_borders)

def add_centered_text(doc, text, size=14, bold=True, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_screenshot_pair(doc, title1, details1, path1, title2, details2, path2):
    """Adds two screenshots with explanations on a single page."""
    # First Screenshot
    p1 = doc.add_paragraph()
    run1 = p1.add_run(title1)
    set_font(run1, size=13, bold=True)
    
    det1 = doc.add_paragraph()
    det1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(det1.add_run(details1), size=11, italic=True)
    
    if os.path.exists(path1):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.add_run().add_picture(path1, width=Inches(5.0))
    doc.add_paragraph()

    # Second Screenshot
    p2 = doc.add_paragraph()
    run2 = p2.add_run(title2)
    set_font(run2, size=13, bold=True)
    
    det2 = doc.add_paragraph()
    det2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(det2.add_run(details2), size=11, italic=True)
    
    if os.path.exists(path2):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(path2, width=Inches(5.0))
    
    doc.add_page_break()

def create_documentation():
    doc = Document()
    set_page_border(doc)
    
    artifacts_dir = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa"
    gls_logo = os.path.join(artifacts_dir, "gls_university_logo_placeholder_1773730030624.png")
    naac_logo = os.path.join(artifacts_dir, "naac_a_plus_grade_logo_placeholder_1773730051169.png")
    
    login_ss = os.path.join(artifacts_dir, "login_page_screenshot_1773730549777.png")
    dashboard_ss = os.path.join(artifacts_dir, "chat_dashboard_1773730844704.png")
    response_ss = os.path.join(artifacts_dir, "ai_response_1773730919335.png")
    users_ss = os.path.join(artifacts_dir, "admin_dashboard_users_1773730983363.png")
    feedback_ss = os.path.join(artifacts_dir, "admin_dashboard_feedback_1773730989927.png")
    analytics_ss = os.path.join(artifacts_dir, "analytics_dashboard_1773731023259.png")

    # --- PAGE 1: COVER PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(gls_logo): p.add_run().add_picture(gls_logo, width=Inches(2.8))
    if os.path.exists(naac_logo): 
        p.add_run("    ")
        p.add_run().add_picture(naac_logo, width=Inches(1.6))
    doc.add_paragraph("\n")
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=18)
    add_centered_text(doc, "Integrated Master of Computer Applications [iMSc.IT] Programme", size=14)
    add_centered_text(doc, "Semester VIII", size=14)
    doc.add_paragraph("\n")
    add_centered_text(doc, "Project Report for", size=12, bold=False)
    add_centered_text(doc, "221601801 CROSS PLATFORM MOBILE\nAPP DEVELOPMENT", size=16)
    doc.add_paragraph("\n")
    add_centered_text(doc, "“Question Answering System”", size=14)
    doc.add_paragraph("\n" * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    set_font(run, size=12, bold=True, underline=True)
    add_centered_text(doc, "Group No: [01]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Yash Modi, [EnrolmentNo]", size=12, bold=False, space_after=2)
    doc.add_page_break()
    
    # --- PAGE 2: CERTIFICATE ---
    add_centered_text(doc, "GLS UNIVERSITY", size=14)
    add_centered_text(doc, "Faculty of Computer Applications & IT\niMSc.IT Programme\nAhmedabad", size=12)
    doc.add_paragraph("\n")
    add_centered_text(doc, "CERTIFICATE", size=16)
    doc.add_paragraph("________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    cert_text = "This is to certify that Yash Modi, Student of Semester- VIII iMSc.IT, FCAIT, GLS University has successfully completed the Project of \"Question Answering System\" as a fulfillment of the study of Semester-VIII."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.add_run(cert_text), size=12)
    doc.add_paragraph("\n\nDate of Submission: ________________\nProject Guide - Name & Sign: ________________")
    doc.add_page_break()
    
    # --- PAGE 3: TABLE OF CONTENTS ---
    add_centered_text(doc, "Table of Contents", size=14)
    items = ["1. Introduction", "   1.1 Project Objective", "   1.2 Purpose", "   1.3 Modules", "   1.4 Workflow & Architecture", "2. Screenshots with Explanation", "   2.1 User-side Interface", "   2.2 Admin-side Management", "3. Conclusion"]
    for it in items:
        p = doc.add_paragraph(it)
        set_font(p.runs[0])
        if "   " in it: p.paragraph_format.left_indent = Inches(0.5)
    doc.add_page_break()
    
    # --- PAGE 4: DETAILED INTRODUCTION ---
    p = doc.add_heading("1. Introduction", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p = doc.add_heading("1.1 Project Objective", level=2)
    set_font(p.runs[0], size=13, bold=True)
    intro_txt = ("The fundamental objective of this project is to implement a Retrieval-Based Question Answering (QA) System. "
                 "Unlike traditional classifiers, this system leverages Retrieval-Augmented Generation (RAG) principles to find "
                 "semantically similar answers from a massive repository of 600,000+ technical QA pairs.")
    set_font(doc.add_paragraph(intro_txt).runs[0])

    p = doc.add_heading("1.2 Purpose", level=2)
    set_font(p.runs[0], size=13, bold=True)
    purpose_txt = ("Information overload makes data retrieval a core challenge. This project solves this by providing a "
                   "unified AI interface where users can ask complex technical questions and receive instantaneous, context-aware "
                   "responses. It demonstrates the integration of machine learning models into a real-world SaaS architecture.")
    set_font(doc.add_paragraph(purpose_txt).runs[0])

    p = doc.add_heading("1.3 Modules", level=2)
    set_font(p.runs[0], size=13, bold=True)
    modules_txt = "• Data Layer: Handles dataset generation and TF-IDF vectorization.\n• Inference Layer: Asynchronous FastAPI server evaluating cosine similarity.\n• Frontend Layer: Vite-React application with real-time UI updates."
    set_font(doc.add_paragraph(modules_txt).runs[0])

    p = doc.add_heading("1.4 Workflow & Architecture", level=2)
    set_font(p.runs[0], size=13, bold=True)
    workflow_txt = ("How the Project Works:\n1. Pre-Processing: Input text is tokenized, lemmatized, and stripped of stopwords using NLTK.\n"
                    "2. Vectorization: The query is converted into a TF-IDF sparse vector.\n"
                    "3. Similarity Search: The system computes Cosine Similarity between the query vector and the FAISS indexed dataset vectors.\n"
                    "4. Generation: If a high-confidence match is found, it is retrieved. Otherwise, a local LLM (Llama3) generates a refined answer based on retrieved context.")
    set_font(doc.add_paragraph(workflow_txt).runs[0])
    doc.add_page_break()

    # --- SCREENSHOTS SECTION (2 PER PAGE) ---
    p = doc.add_heading("2. Screenshots with Explanation", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    # Page 1 of Screenshots
    p_u = doc.add_heading("2.1 User-side Interface", level=2)
    set_font(p_u.runs[0], size=13, bold=True)
    
    add_screenshot_pair(doc, 
        "1. Authentication System", 
        "The login page provides secure access to the AI engine. Users can sign up or log in with their credentials to access personalized chat histories.",
        login_ss,
        "2. Main Chat Dashboard",
        "The dashboard features a 3-pane layout with a sticky sidebar for chat history, a modern chat window, and a system analytics panel.",
        dashboard_ss
    )

    # Page 2 of Screenshots
    add_screenshot_pair(doc,
        "3. AI Response & Confidence Score",
        "Every AI response displays a confidence score, indicating the mathematical similarity of the retrieved answer to the user's question.",
        response_ss,
        "4. Admin Statistics (Analytics)",
        "The statistics page provides a breakdown of system performance, including total records processed, average latency, and domain distribution.",
        analytics_ss
    )

    # Page 3 of Screenshots
    p_a = doc.add_heading("2.2 Admin-side Management", level=2)
    set_font(p_a.runs[0], size=13, bold=True)
    add_screenshot_pair(doc,
        "5. User Management Panel",
        "Admins can monitor all registered users, manage roles (Admin/User), and handle account operations like password resets or deletions.",
        users_ss,
        "6. Feedback Management",
        "The system collects up/down votes on AI answers, allowing admins to review where factual accuracy might need improvement.",
        feedback_ss
    )

    doc.save("Project_Report_Detailed_V4.docx")
    print("Detailed report (V4) generated.")

if __name__ == "__main__":
    create_documentation()
