from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def set_page_border(doc):
    """Adds a page border to all sections in the document."""
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '12') # 1.5 pt
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)
        sect_pr.append(pg_borders)

def add_centered_text(doc, text, size=14, bold=True, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_screenshot(doc, title, path):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, size=12, italic=True)
    if os.path.exists(path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(path, width=Inches(6.0))
    else:
        p2 = doc.add_paragraph(f"[SCREENSHOT AT {path} NOT FOUND]")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

def create_documentation():
    doc = Document()
    set_page_border(doc)
    
    # Paths
    artifacts_dir = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa"
    gls_logo = os.path.join(artifacts_dir, "gls_university_logo_placeholder_1773730030624.png")
    naac_logo = os.path.join(artifacts_dir, "naac_a_plus_grade_logo_placeholder_1773730051169.png")
    
    # Screenshots
    login_ss = os.path.join(artifacts_dir, "login_page_screenshot_1773730549777.png")
    dashboard_ss = os.path.join(artifacts_dir, "chat_dashboard_1773730844704.png")
    response_ss = os.path.join(artifacts_dir, "ai_response_1773730919335.png")
    users_ss = os.path.join(artifacts_dir, "admin_dashboard_users_1773730983363.png")
    feedback_ss = os.path.join(artifacts_dir, "admin_dashboard_feedback_1773730989927.png")
    analytics_ss = os.path.join(artifacts_dir, "analytics_dashboard_1773731023259.png")

    # --- PAGE 1: COVER PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(gls_logo): p.add_run().add_picture(gls_logo, width=Inches(3.0))
    if os.path.exists(naac_logo): 
        p.add_run("    ")
        p.add_run().add_picture(naac_logo, width=Inches(1.8))
    
    doc.add_paragraph("\n")
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=18)
    add_centered_text(doc, "Integrated Master of Computer Applications [iMSc.IT] Programme", size=14)
    add_centered_text(doc, "Semester VIII", size=14)
    doc.add_paragraph("\n")
    add_centered_text(doc, "Project Report for", size=12, bold=False)
    add_centered_text(doc, "221601801 CROSS PLATFORM MOBILE\nAPP DEVELOPMENT", size=16)
    doc.add_paragraph("\n")
    add_centered_text(doc, "“Question Answering System”", size=14)
    doc.add_paragraph("\n" * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    set_font(run, size=12, bold=True, underline=True)
    add_centered_text(doc, "Group No: [01]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Yash Modi, [EnrolmentNo]", size=12, bold=False, space_after=2)
    doc.add_page_break()
    
    # --- PAGE 2: CERTIFICATE ---
    add_centered_text(doc, "GLS UNIVERSITY", size=14)
    add_centered_text(doc, "Faculty of Computer Applications & IT\niMSc.IT Programme\nAhmedabad", size=12)
    doc.add_paragraph("\n")
    add_centered_text(doc, "CERTIFICATE", size=16)
    doc.add_paragraph("________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    cert_text = "This is to certify that Yash Modi, Student of Semester- VIII iMSc.IT, FCAIT, GLS University has successfully completed the Project of \"Question Answering System\" as a fulfillment of the study of Semester-VIII."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.add_run(cert_text), size=12)
    doc.add_paragraph("\n\nDate of Submission: ________________\nProject Guide - Name & Sign: ________________")
    doc.add_page_break()
    
    # --- PAGE 3: TABLE OF CONTENTS ---
    add_centered_text(doc, "Table of Contents", size=14)
    items = ["1. Introduction", "   1.1 Project Objective", "   1.2 Purpose", "   1.3 Modules", "2. Screenshots", "   2.1 User-side", "   2.2 Admin-side"]
    for it in items:
        p = doc.add_paragraph(it)
        set_font(p.runs[0])
        if "   " in it: p.paragraph_format.left_indent = Inches(0.5)
    doc.add_page_break()
    
    # --- PAGE 4: INTRODUCTION ---
    p = doc.add_heading("1. Introduction", level=1)
    set_font(p.runs[0], size=14, bold=True)
    p = doc.add_heading("1.1 Project Objective", level=2)
    set_font(p.runs[0], size=13, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.add_run("The objective is to provide a semantic search engine using TF-IDF and Cosine Similarity to answer technical questions with high accuracy."))
    doc.add_page_break()

    # --- PAGE 5: USER SIDE ---
    p = doc.add_heading("2. Screenshots", level=1)
    set_font(p.runs[0], size=14, bold=True)
    p = doc.add_heading("2.1 User-side", level=2)
    set_font(p.runs[0], size=13, bold=True)
    
    add_screenshot(doc, "Login Interface - [Widget Name: AuthForm]", login_ss)
    doc.add_page_break()
    add_screenshot(doc, "Chat Dashboard - [Widget Name: ChatContainer]", dashboard_ss)
    doc.add_page_break()
    add_screenshot(doc, "AI Response with Confidence Score - [Widget Name: ResponseBubble]", response_ss)
    doc.add_page_break()
    
    # --- PAGE 8: ADMIN SIDE ---
    p = doc.add_heading("2.2 Admin-side", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    add_screenshot(doc, "User Management - [Widget Name: UserTable]", users_ss)
    doc.add_page_break()
    add_screenshot(doc, "User Feedback Monitoring - [Widget Name: FeedbackList]", feedback_ss)
    doc.add_page_break()
    add_screenshot(doc, "System Analytics Dashboard - [Widget Name: StatsCharts]", analytics_ss)

    doc.save("Project_Report_Final.docx")
    print("Final report generated.")

if __name__ == "__main__":
    create_documentation()
