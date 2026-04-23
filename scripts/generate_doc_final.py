from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def add_centered_text(doc, text, size=14, bold=True, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def create_documentation():
    doc = Document()
    
    # --- PAGE 1: COVER PAGE ---
    # Placeholder for Logo
    add_centered_text(doc, "[INSERT GLS UNIVERSITY LOGO HERE]", size=10, bold=False)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=18)
    add_centered_text(doc, "Integrated Master of Computer Applications [iMSc.IT] Programme", size=14)
    add_centered_text(doc, "Semester VIII", size=14)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "Project Report for", size=12, bold=False)
    add_centered_text(doc, "Retrieval-Based Question Answering System", size=16)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "“AI Knowledge Assistant”", size=14)
    doc.add_paragraph("\n" * 4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    set_font(run, size=12, bold=True)
    run.underline = True
    
    add_centered_text(doc, "Group No: [Your Group Number]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Student Name 1, [Enrolment No]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Student Name 2, [Enrolment No]", size=12, bold=False, space_after=2)
    
    doc.add_page_break()
    
    # --- PAGE 2: CERTIFICATE ---
    add_centered_text(doc, "GLS UNIVERSITY", size=14)
    add_centered_text(doc, "Faculty of Computer Applications & IT", size=12)
    add_centered_text(doc, "iMSc.IT Programme", size=12)
    add_centered_text(doc, "Ahmedabad", size=12)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "CERTIFICATE", size=16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bar = p.add_run("________________________________________")
    set_font(bar, size=12, bold=True)
    
    doc.add_paragraph("\n")
    cert_text = (
        "This is to certify that:\n"
        "1) [Student Name 1]\n"
        "2) [Student Name 2]\n\n"
        "Students of Semester- VIII iMSc.IT, FCAIT, GLS University have successfully completed the Project of "
        "\"Retrieval-Based Question Answering System\" as a fulfillment of the study of Semester-VIII, "
        "Integrated Master of Computer Applications [iMSc.IT]."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(cert_text)
    set_font(run, size=12)
    
    doc.add_paragraph("\n" * 2)
    p = doc.add_paragraph()
    run = p.add_run("Date of Submission: ________________")
    set_font(run, size=12)
    
    p = doc.add_paragraph()
    run = p.add_run("Project Guide - Name & Sign: ________________")
    set_font(run, size=12)
    
    doc.add_page_break()
    
    # --- PAGE 3: TABLE OF CONTENTS ---
    add_centered_text(doc, "Table of Contents", size=14)
    toc_items = [
        "1. Introduction",
        "   1.1 Project Objective",
        "   1.2 Purpose",
        "   1.3 Modules",
        "2. System Architecture",
        "3. Screenshots",
        "   3.1 User-side",
        "   3.2 Admin-side",
        "4. Conclusion"
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, size=12)
        p.paragraph_format.left_indent = Inches(0.5) if any(x in item for x in ["1.1", "1.2", "1.3", "3.1", "3.2"]) else Inches(0)

    doc.add_page_break()
    
    # --- PAGE 4: INTRODUCTION ---
    p = doc.add_heading("1. Introduction", level=1)
    run = p.runs[0]
    set_font(run, size=14, bold=True)
    
    p = doc.add_heading("1.1 Project Objective", level=2)
    set_font(p.runs[0], size=13, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "The primary objective of this project is to develop a high-performance Question Answering system "
        "capable of retrieving accurate information from a large-scale synthetic dataset of over 600,000 pairs. "
        "The system aims to bridge the gap between simple keyword searching and intelligent semantic retrieval."
    )
    set_font(p.add_run(text))
    
    p = doc.add_heading("1.2 Purpose", level=2)
    set_font(p.runs[0], size=13, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "With the exponential growth of digital data, retrieving specific information quickly is critical. "
        "This project serves as a demonstration of NLP techniques like TF-IDF Vectorization and Cosine Similarity "
        "integrated into a modern web architecture, providing users with sub-50ms response times for complex queries."
    )
    set_font(p.add_run(text))
    
    p = doc.add_heading("1.3 Modules", level=2)
    set_font(p.runs[0], size=13, bold=True)
    modules = [
        ("Data Generation & Cleaning", "Scripts to synthesize 600k+ QA pairs and normalize them using NLTK lemmatization."),
        ("Retrieval Engine", "A TF-IDF based matching system that utilizes sparse matrices for memory-efficient similarity scoring."),
        ("FastAPI Backend", "An asynchronous Python API that handles real-time inference and serves the ML model results."),
        ("React Frontend", "A modern, responsive dashboard with glassmorphism design and real-time confidence score visualization."),
        ("Admin Dashboard", "A secure interface for monitoring system stats, user feedback, and managing users.")
    ]
    for m_title, m_desc in modules:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_title = p.add_run(f"• {m_title}: ")
        set_font(run_title, bold=True)
        run_desc = p.add_run(m_desc)
        set_font(run_desc)

    doc.add_page_break()
    
    # --- PAGE 5: SCREENSHOTS ---
    p = doc.add_heading("3. Screenshots", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p = doc.add_heading("3.1 User-side", level=2)
    set_font(p.runs[0], size=13, bold=True)
    
    placeholders = [
        "Login / Register Page - [Widget Name: AuthForm]",
        "Chat Interface Dashboard - [Widget Name: ChatContainer]",
        "AI Response with Confidence Score - [Widget Name: ResponseBubble]",
        "History Sidebar - [Widget Name: SidebarSessions]"
    ]
    for ph in placeholders:
        p = doc.add_paragraph()
        run = p.add_run(ph)
        set_font(run, italic=True)
        p = doc.add_paragraph("[INSERT SCREENSHOT HERE]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

    p = doc.add_heading("3.2 Admin-side", level=2)
    set_font(p.runs[0], size=13, bold=True)
    
    admin_placeholders = [
        "Admin Statistics Overview - [Widget Name: StatsDashboard]",
        "Feedback Management List - [Widget Name: FeedbackTable]",
        "User Management Control - [Widget Name: UserManagement]"
    ]
    for ph in admin_placeholders:
        p = doc.add_paragraph()
        run = p.add_run(ph)
        set_font(run, italic=True)
        p = doc.add_paragraph("[INSERT SCREENSHOT HERE]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

    # Save the document
    output_path = "Project_Documentation.docx"
    doc.save(output_path)
    print(f"Documentation saved as {output_path}")

if __name__ == "__main__":
    create_documentation()
