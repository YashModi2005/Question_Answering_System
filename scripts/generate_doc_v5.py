from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def set_page_border(doc):
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)
        sect_pr.append(pg_borders)

def add_centered_text(doc, text, size=14, bold=True, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_screenshot_pair(doc, title1, details1, path1, title2, details2, path2):
    # First Screenshot
    p1 = doc.add_paragraph()
    run1 = p1.add_run(title1)
    set_font(run1, size=13, bold=True)
    
    det1 = doc.add_paragraph()
    det1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(det1.add_run(details1), size=11, italic=True)
    
    if os.path.exists(path1):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.add_run().add_picture(path1, width=Inches(5.0))
    doc.add_paragraph()

    # Second Screenshot
    p2 = doc.add_paragraph()
    run2 = p2.add_run(title2)
    set_font(run2, size=13, bold=True)
    
    det2 = doc.add_paragraph()
    det2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(det2.add_run(details2), size=11, italic=True)
    
    if os.path.exists(path2):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(path2, width=Inches(5.0))
    
    doc.add_page_break()

def create_documentation():
    doc = Document()
    set_page_border(doc)
    
    artifacts_dir = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa"
    
    login_ss = os.path.join(artifacts_dir, "login_page_screenshot_1773730549777.png")
    dashboard_ss = os.path.join(artifacts_dir, "chat_dashboard_1773730844704.png")
    response_ss = os.path.join(artifacts_dir, "ai_response_1773730919335.png")
    users_ss = os.path.join(artifacts_dir, "admin_dashboard_users_1773730983363.png")
    feedback_ss = os.path.join(artifacts_dir, "admin_dashboard_feedback_1773730989927.png")
    analytics_ss = os.path.join(artifacts_dir, "analytics_dashboard_1773731023259.png")

    # --- PAGE 1: COVER PAGE (Logos Removed) ---
    doc.add_paragraph("\n" * 4)
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=18)
    add_centered_text(doc, "Integrated Master of Computer Applications [iMSc.IT] Programme", size=14)
    add_centered_text(doc, "Semester VIII", size=14)
    doc.add_paragraph("\n" * 2)
    add_centered_text(doc, "Project Report for", size=12, bold=False)
    add_centered_text(doc, "221601801 CROSS PLATFORM MOBILE\nAPP DEVELOPMENT", size=16)
    doc.add_paragraph("\n")
    add_centered_text(doc, "“Question Answering System”", size=15)
    doc.add_paragraph("\n" * 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    set_font(run, size=12, bold=True, underline=True)
    add_centered_text(doc, "Group No: [01]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Yash Modi, [EnrolmentNo]", size=12, bold=False, space_after=2)
    doc.add_page_break()
    
    # --- PAGE 2: CERTIFICATE ---
    add_centered_text(doc, "GLS UNIVERSITY", size=14)
    add_centered_text(doc, "Faculty of Computer Applications & IT\niMSc.IT Programme\nAhmedabad", size=12)
    doc.add_paragraph("\n")
    add_centered_text(doc, "CERTIFICATE", size=16)
    doc.add_paragraph("________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    cert_text = "This is to certify that Yash Modi, Student of Semester- VIII iMSc.IT, FCAIT, GLS University has successfully completed the Project of \"Question Answering System\" as a fulfillment of the study of Semester-VIII."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.add_run(cert_text), size=12)
    doc.add_paragraph("\n\nDate of Submission: ________________\nProject Guide - Name & Sign: ________________")
    doc.add_page_break()
    
    # --- PAGE 3: TABLE OF CONTENTS ---
    add_centered_text(doc, "Table of Contents", size=14)
    items = ["1. Introduction", "   1.1 Project Objective", "   1.2 Purpose", "   1.3 Modules", "   1.4 Workflow & Architecture", "2. Screenshots with Explanation", "   2.1 User-side Interface", "   2.2 Admin-side Management", "3. Conclusion"]
    for it in items:
        p = doc.add_paragraph(it)
        set_font(p.runs[0])
        if "   " in it: p.paragraph_format.left_indent = Inches(0.5)
    doc.add_page_break()
    
    # --- PAGE 4: DETAILED INTRODUCTION ---
    p = doc.add_heading("1. Introduction", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p = doc.add_heading("1.1 Project Objective", level=2)
    set_font(p.runs[0], size=13, bold=True)
    intro_txt = ("The primary objective of this project is to implement a robust and scalable Retrieval-Based Question Answering (QA) System. "
                 "Unlike generative-only models that may hallucinate, this system prioritizes factual accuracy by utilizing a curated knowledge base. "
                 "By leveraging Retrieval-Augmented Generation (RAG) principles, the system identifies the most semantically relevant answers from a massive repository of 600,000+ technical QA pairs, "
                 "ensuring that the information provided is both verified and precise. The core goal is to bridge the gap between simple keyword search engines and high-level AI assistants.")
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p_intro.add_run(intro_txt), size=12)

    p = doc.add_heading("1.2 Purpose", level=2)
    set_font(p.runs[0], size=13, bold=True)
    purpose_txt = ("In the modern digital era, the volume of technical documentation and data is growing exponentially, making rapid information retrieval a major bottleneck. "
                   "The purpose of this project is to demonstrate a highly efficient, sub-50ms latency solution for information discovery. "
                   "It serves as a professional-grade demonstration of how Natural Language Processing (NLP) techniques, such as TF-IDF (Term Frequency-Inverse Document Frequency) "
                   "and Cosine Similarity, can be integrated into a functional SaaS (Software as a Service) architecture. The project aim is to empower developers and students "
                   "with an instant AI knowledge assistant for technical queries.")
    p_purp = doc.add_paragraph()
    p_purp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p_purp.add_run(purpose_txt), size=12)

    p = doc.add_heading("1.3 Modules", level=2)
    set_font(p.runs[0], size=13, bold=True)
    m1 = "• Data Management & Pre-processing: Handles the synthesis of large-scale datasets, cleaning noise, and applying NLP normalization techniques like lemmatization."
    m2 = "• Retrieval & Vectorization Engine: Utilizes sparse matrix storage and TF-IDF vectorization to convert human language into searchable mathematical vectors."
    m3 = "• FastAPI Backend Infrastructure: A high-performance, asynchronous Python web server that manages requests, evaluates model similarity scores, and interfaces with the LLM."
    m4 = "• React Dashboard (Frontend): A premium, responsive user interface featuring glassmorphism design, real-time message streaming, and confidence score visualization."
    m5 = "• Admin Control Center: A secure auditing module for monitoring system usage statistics, tracking user feedback, and managing global system configurations."
    for m in [m1, m2, m3, m4, m5]:
        p_m = doc.add_paragraph(m)
        p_m.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_font(p_m.runs[0], size=12)

    p = doc.add_heading("1.4 Workflow & Architecture", level=2)
    set_font(p.runs[0], size=13, bold=True)
    workflow_txt = ("How the System Functions (Technical Pipeline):\n"
                    "1. Linguistic Pre-Processing: When a user submits a question, the text undergoes tokenization (breaking into words), lemmatization (reducing words to their root form), and stopword removal to focus on key semantic terms.\n"
                    "2. Query Vectorization: The processed question is transformed into a TF-IDF vector. This vector represents the relative importance of words within the context of the entire technical dataset.\n"
                    "3. FAISS Similarity Search: The system performs a high-speed vector search (utilizing FAISS indices) to calculate the Cosine Similarity between the user's query and the pre-computed vectors of the 600,000+ stored answers.\n"
                    "4. Hybrid Generation Mechanism: If the similarity score exceeds a predefined threshold (High Confidence), the exact answer is retrieved instantly. If confidence is lower, the system feeds the top 'k' retrieved contexts into a local Llama3 LLM to synthesize a refined, contextually accurate response.")
    p_work = doc.add_paragraph(workflow_txt)
    p_work.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p_work.runs[0], size=12)
    doc.add_page_break()

    # --- SCREENSHOTS SECTION (2 PER PAGE) ---
    p = doc.add_heading("2. Screenshots with Explanation", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p_u = doc.add_heading("2.1 User-side Interface", level=2)
    set_font(p_u.runs[0], size=13, bold=True)
    
    add_screenshot_pair(doc, 
        "1. Secure Authentication Gateway", 
        "The login system ensures that all technical inquiries are tracked and personalized. It uses hashed passwords and role-based access control to distinguish between standard users and system administrators.",
        login_ss,
        "2. Knowledge Assistant Dashboard",
        "The main interface is designed with a focus on modern UX, featuring a sidebar for session management and a central chat area with real-time status indicators for the AI engine.",
        dashboard_ss
    )

    add_screenshot_pair(doc,
        "3. Semantic Response Visualization",
        "This screenshot highlights the retrieval accuracy. Notice the confidence score and the matched dataset question, providing transparency into how the AI arrived at the specific technical answer.",
        response_ss,
        "4. System Performance Analytics",
        "The analytics module tracks the health of the system, displaying the total record count (600k+), average inference latency, and the primary technical domains covered by the knowledge base.",
        analytics_ss
    )

    p_a = doc.add_heading("2.2 Admin-side Management", level=2)
    set_font(p_a.runs[0], size=13, bold=True)
    add_screenshot_pair(doc,
        "5. Administrative User Audit",
        "Admins have a birds-eye view of all registered members. This module allows for password resets, role modifications, and account deletions to maintain system integrity.",
        users_ss,
        "6. Quality Control (Feedback Loop)",
        "The feedback system captures user satisfaction through binary Up/Down votes. This data is critical for fine-tuning the retrieval engine and identifying gaps in the technical knowledge base.",
        feedback_ss
    )

    doc.save("Project_Report_Refined_V5.docx")
    print("Elaborated report (V5) generated without logos.")

if __name__ == "__main__":
    create_documentation()
