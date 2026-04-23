from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def set_page_border(doc):
    """Adds a page border to all sections in the document."""
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '12') # 1.5 pt
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)
        sect_pr.append(pg_borders)

def add_centered_text(doc, text, size=14, bold=True, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def create_documentation():
    doc = Document()
    set_page_border(doc)
    
    # Paths to generated logos
    gls_logo = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa\gls_university_logo_placeholder_1773730030624.png"
    naac_logo = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa\naac_a_plus_grade_logo_placeholder_1773730051169.png"
    
    # Paths to screenshots (Placeholders if not captured yet)
    # I'll use placeholders for now, but the script will look for them.
    screenshots_dir = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa"
    login_ss = os.path.join(screenshots_dir, "login_page_screenshot_1773730549777.png")
    dashboard_ss = os.path.join(screenshots_dir, "chat_dashboard_yash_1773730603537.png")
    # Add more as they come...

    # --- PAGE 1: COVER PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if os.path.exists(gls_logo):
        p.add_run().add_picture(gls_logo, width=Inches(2.5))
    if os.path.exists(naac_logo):
        p.add_run("    ") # Spacer
        p.add_run().add_picture(naac_logo, width=Inches(1.5))
    
    doc.add_paragraph("\n")
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=18)
    add_centered_text(doc, "Integrated Master of Computer Applications [iMSc.IT] Programme", size=14)
    add_centered_text(doc, "Semester VIII", size=14)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "Project Report for", size=12, bold=False)
    add_centered_text(doc, "221601801 CROSS PLATFORM MOBILE\nAPP DEVELOPMENT", size=16)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "“Retrieval-Based Question Answering System”", size=14)
    doc.add_paragraph("\n" * 3)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    set_font(run, size=12, bold=True, underline=True)
    
    add_centered_text(doc, "Group No: [Your Group Number]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Yash Modi, [Enrolment No]", size=12, bold=False, space_after=2)
    
    doc.add_page_break()
    
    # --- PAGE 2: CERTIFICATE ---
    add_centered_text(doc, "GLS UNIVERSITY", size=14)
    add_centered_text(doc, "Faculty of Computer Applications & IT", size=12)
    add_centered_text(doc, "iMSc.IT Programme", size=12)
    add_centered_text(doc, "Ahmedabad", size=12)
    doc.add_paragraph("\n")
    
    add_centered_text(doc, "CERTIFICATE", size=16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bar = p.add_run("________________________________________")
    set_font(bar, size=12, bold=True)
    
    doc.add_paragraph("\n")
    cert_text = (
        "This is to certify that:\n"
        "1) Yash Modi\n"
        "2) [Student Name 2]\n\n"
        "Students of Semester- VIII iMSc.IT, FCAIT, GLS University have successfully completed the Project of "
        "\"Retrieval-Based Question Answering System\" as a fulfillment of the study of Semester-VIII, "
        "Integrated Master of Computer Applications [iMSc.IT]."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(cert_text)
    set_font(run, size=12)
    
    doc.add_paragraph("\n" * 2)
    p = doc.add_paragraph()
    run = p.add_run("Date of Submission: ________________")
    set_font(run, size=12)
    
    p = doc.add_paragraph()
    run = p.add_run("Project Guide - Name & Sign: ________________")
    set_font(run, size=12)
    
    doc.add_page_break()
    
    # --- PAGE 3: TABLE OF CONTENTS ---
    # (Same as before)
    # ... code omitted for brevity in thought, but I will include it in the write_to_file
    add_centered_text(doc, "Table of Contents", size=14)
    toc_items = ["1. Introduction", "   1.1 Project Objective", "   1.2 Purpose", "   1.3 Modules", "2. System Architecture", "3. Screenshots", "   3.1 User-side", "   3.2 Admin-side", "4. Conclusion"]
    for item in toc_items:
        p = doc.add_paragraph(item)
        set_font(p.runs[0], size=12)
        if "   " in item: p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()
    
    # --- PAGE 4: INTRODUCTION ---
    p = doc.add_heading("1. Introduction", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p = doc.add_heading("1.1 Project Objective", level=2)
    set_font(p.runs[0], size=13, bold=True)
    p = doc.add_paragraph("High-performance semantic QA retrieval from 600,000+ technical pairs.")
    set_font(p.runs[0], size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_heading("1.2 Purpose", level=2)
    set_font(p.runs[0], size=13, bold=True)
    p = doc.add_paragraph("Demonstrating advanced NLP (TF-IDF, Cosine Similarity) in a full-stack SaaS dashboard.")
    set_font(p.runs[0], size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

    # --- PAGE 5: SCREENSHOTS ---
    p = doc.add_heading("3. Screenshots", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    def add_screenshot(title, path):
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_font(run, italic=True)
        if os.path.exists(path):
            doc.add_paragraph().add_run().add_picture(path, width=Inches(5.5))
        else:
            p2 = doc.add_paragraph("[SCREENSHOT NOT FOUND - PLACEHOLDER]")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

    add_screenshot("Login Page - [Widget: AuthForm]", login_ss)
    add_screenshot("Main Dashboard - [Widget: ChatContainer]", dashboard_ss)
    
    # Save the document
    output_path = "Project_Documentation_V2.docx"
    doc.save(output_path)
    print(f"Documentation saved as {output_path}")

if __name__ == "__main__":
    create_documentation()
