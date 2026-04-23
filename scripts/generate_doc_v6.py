from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def set_page_border(doc):
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)
        sect_pr.append(pg_borders)

def add_centered_text(doc, text, size=14, bold=True, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_screenshot_pair(doc, title1, details1, path1, title2, details2, path2):
    # First Screenshot
    p1 = doc.add_paragraph()
    run1 = p1.add_run(title1)
    set_font(run1, size=13, bold=True)
    
    det1 = doc.add_paragraph()
    det1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(det1.add_run(details1), size=11, italic=True)
    
    if os.path.exists(path1):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.add_run().add_picture(path1, width=Inches(5.0))
    doc.add_paragraph()

    # Second Screenshot
    p2 = doc.add_paragraph()
    run2 = p2.add_run(title2)
    set_font(run2, size=13, bold=True)
    
    det2 = doc.add_paragraph()
    det2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(det2.add_run(details2), size=11, italic=True)
    
    if os.path.exists(path2):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(path2, width=Inches(5.0))
    
    doc.add_page_break()

def create_documentation():
    doc = Document()
    set_page_border(doc)
    
    artifacts_dir = r"C:\Users\yashu\.gemini\antigravity\brain\44b160c9-16a1-4b53-94a2-8a357ba0bcaa"
    
    login_ss = os.path.join(artifacts_dir, "login_page_screenshot_1773730549777.png")
    dashboard_ss = os.path.join(artifacts_dir, "chat_dashboard_1773730844704.png")
    response_ss = os.path.join(artifacts_dir, "ai_response_1773730919335.png")
    users_ss = os.path.join(artifacts_dir, "admin_dashboard_users_1773730983363.png")
    feedback_ss = os.path.join(artifacts_dir, "admin_dashboard_feedback_1773730989927.png")
    analytics_ss = os.path.join(artifacts_dir, "analytics_dashboard_1773731023259.png")

    # --- PAGE 1: COVER PAGE (Updated with Yash & Nitya) ---
    doc.add_paragraph("\n" * 4)
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=18)
    add_centered_text(doc, "Integrated Master of Computer Applications [iMSc.IT] Programme", size=14)
    add_centered_text(doc, "Semester VIII", size=14)
    doc.add_paragraph("\n" * 2)
    add_centered_text(doc, "Project Report for", size=12, bold=False)
    add_centered_text(doc, "221601804 Advanced Machine Learning and\nDeep Learning", size=16)
    doc.add_paragraph("\n")
    add_centered_text(doc, "“Question Answering System”", size=15)
    doc.add_paragraph("\n" * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    set_font(run, size=12, bold=True, underline=True)
    add_centered_text(doc, "Group No: [42]", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Yash Modi, 202201619010159", size=12, bold=False, space_after=2)
    add_centered_text(doc, "Nitya Modi, 202201619010046", size=12, bold=False, space_after=2)
    doc.add_page_break()
    
    # --- PAGE 2: CERTIFICATE ---
    add_centered_text(doc, "GLS UNIVERSITY", size=14)
    add_centered_text(doc, "Faculty of Computer Applications & IT\niMSc.IT Programme\nAhmedabad", size=12)
    doc.add_paragraph("\n")
    add_centered_text(doc, "CERTIFICATE", size=16)
    doc.add_paragraph("________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    cert_text = "This is to certify that Yash Modi and Nitya Modi, Students of Semester- VIII iMSc.IT, FCAIT, GLS University have successfully completed the Project of \"Question Answering System\" as a fulfillment of the study of Semester-VIII."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.add_run(cert_text), size=12)
    doc.add_paragraph("\n\nDate of Submission: ________________\nProject Guide - Name & Sign: ________________")
    doc.add_page_break()
    
    # --- PAGE 3: TABLE OF CONTENTS ---
    add_centered_text(doc, "Table of Contents", size=14)
    items = ["1. Introduction", "   1.1 Project Goal", "   1.2 Purpose", "   1.3 System Parts", "   1.4 How it Works", "2. Screenshots and Details", "   2.1 User View", "   2.2 Admin View", "3. Conclusion"]
    for it in items:
        p = doc.add_paragraph(it)
        set_font(p.runs[0])
        if "   " in it: p.paragraph_format.left_indent = Inches(0.5)
    doc.add_page_break()
    
    # --- PAGE 4: SIMPLE INTRODUCTION ---
    p = doc.add_heading("1. Introduction", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p = doc.add_heading("1.1 Project Goal", level=2)
    set_font(p.runs[0], size=13, bold=True)
    intro_txt = ("The main goal of this project is to build an easy-to-use system that can answer technical questions. "
                 "Instead of just searching for words, it looks for the meaning behind the question to find the best answer from a large list of 600,000+ technical questions and answers.")
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p_intro.add_run(intro_txt), size=12)

    p = doc.add_heading("1.2 Purpose", level=2)
    set_font(p.runs[0], size=13, bold=True)
    purpose_txt = ("Sometimes it is very hard to find the right information in a lot of data. This project helps people by providing "
                   "an AI assistant that gives answers quickly. It uses smart math techniques to match questions with the right answers instantly.")
    p_purp = doc.add_paragraph()
    p_purp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p_purp.add_run(purpose_txt), size=12)

    p = doc.add_heading("1.3 System Parts (Modules)", level=2)
    set_font(p.runs[0], size=13, bold=True)
    m1 = "• Data Part: This part cleans the data and prepares it for the search engine."
    m2 = "• Search Part: This part converts human words into numbers so the computer can find matches."
    m3 = "• Server Part: This is the brain of the project that handles user questions and finds the answers."
    m4 = "• User Part: This is what the user sees—a simple and beautiful screen to type questions and see answers."
    m5 = "• Admin Part: This is for the manager to see how many people are using the system and what they think about it."
    for m in [m1, m2, m3, m4, m5]:
        p_m = doc.add_paragraph(m)
        p_m.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_font(p_m.runs[0], size=12)

    p = doc.add_heading("1.4 How it Works", level=2)
    set_font(p.runs[0], size=13, bold=True)
    workflow_txt = ("Steps the system takes:\n"
                    "1. Cleaning: The system cleans the user's question by removing extra words.\n"
                    "2. Converting: It turns the question into a special format that the computer understands.\n"
                    "3. Finding: It looks through the large dataset to find the most similar answer.\n"
                    "4. Showing: If it finds a good match, it shows it to the user. If not, a smart AI helps to create a fresh answer.")
    p_work = doc.add_paragraph(workflow_txt)
    p_work.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p_work.runs[0], size=12)
    doc.add_page_break()

    # --- SCREENSHOTS SECTION (2 PER PAGE) ---
    p = doc.add_heading("2. Screenshots and Details", level=1)
    set_font(p.runs[0], size=14, bold=True)
    
    p_u = doc.add_heading("2.1 User View", level=2)
    set_font(p_u.runs[0], size=13, bold=True)
    
    add_screenshot_pair(doc, 
        "1. Login Screen", 
        "This is where users enter their name and password to start using the system safely.",
        login_ss,
        "2. Chat Screen",
        "This is the main screen where users can type their questions and see their past chats.",
        dashboard_ss
    )

    add_screenshot_pair(doc,
        "3. AI Answer and Match Score",
        "This screen shows the answer provided by the AI. It also shows a 'Score' which tells us how sure the AI is about the answer.",
        response_ss,
        "4. System Dashboard",
        "This shows how much data is in the system and how fast it is working.",
        analytics_ss
    )

    p_a = doc.add_heading("2.2 Admin View", level=2)
    set_font(p_a.runs[0], size=13, bold=True)
    add_screenshot_pair(doc,
        "5. User Management",
        "The admin can see all the users who have joined the system here.",
        users_ss,
        "6. User Feedback",
        "The admin can see if users liked the answers or not, so we can make the system better.",
        feedback_ss
    )

    doc.save("Project_Report_Final_Simplified_V6.docx")
    print("Simplified report (V6) generated.")

if __name__ == "__main__":
    create_documentation()
